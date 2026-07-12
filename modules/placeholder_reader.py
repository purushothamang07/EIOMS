import re
from docx import Document
from openpyxl import load_workbook


def get_word_placeholders(doc_path):

    doc = Document(doc_path)

    placeholders = set()

    pattern = r"{{(.*?)}}"

    for para in doc.paragraphs:
        matches = re.findall(pattern, para.text)
        for m in matches:
            placeholders.add(m.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(pattern, cell.text)
                for m in matches:
                    placeholders.add(m.strip())

    return placeholders


def get_excel_placeholders(file_path):

    wb = load_workbook(file_path)

    placeholders = set()

    pattern = r"{{(.*?)}}"

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    matches = re.findall(pattern, cell.value)
                    for m in matches:
                        placeholders.add(m.strip())

    return placeholders


def get_placeholders(file_path):

    print("========== DEBUG ==========")
    print("Using NEW placeholder_reader")
    print("File:", file_path)
    print("===========================")

    if file_path.lower().endswith(".docx"):
        return sorted(get_word_placeholders(file_path))

    elif file_path.lower().endswith(".xlsx"):
        return sorted(get_excel_placeholders(file_path))

    else:
        print("Unsupported file type")
        return []